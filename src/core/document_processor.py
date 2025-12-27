import os
import re
from typing import List, Dict, Optional
from pathlib import Path
import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

# Optional imports with fallbacks
try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    import docx2txt
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_metadata(self, file_path: str, content: str) -> Dict:
        """Extract enhanced metadata for filtering"""
        file_name = Path(file_path).name.lower()
        metadata = {
            "source": file_path,
            "file_name": Path(file_path).name,
            "file_type": Path(file_path).suffix.lower(),
        }
        
        # Detect document type - files in company folders are placement papers
        if any(company in str(Path(file_path).parent).lower() for company in ['cocubes', 'mphasis', 'valuelabs', 'zenq']):
            metadata["document_type"] = "placement_paper"
        elif any(term in file_name for term in ['mock', 'test', 'exam', 'quiz']):
            metadata["document_type"] = "mock_test"
        else:
            metadata["document_type"] = "learning_material"
        
        # Extract company names from folder path and filename
        folder_path = str(Path(file_path).parent).lower()
        companies = ['cocubes', 'mphasis', 'valuelabs', 'zenq', 'google', 'microsoft', 'amazon', 'apple', 'meta', 'netflix', 'uber', 'airbnb']
        for company in companies:
            if company in folder_path or company in file_name:
                metadata["company"] = company.title()
                break
        
        # Extract year
        year_match = re.search(r'20\d{2}', file_name)
        if year_match:
            metadata["year"] = year_match.group()
        
        # Extract subject/topic from placement papers
        subjects = {
            'quant': 'Quantitative Aptitude', 'logical': 'Logical Reasoning', 'english': 'English',
            'computer': 'Computer Fundamentals', 'reasoning': 'Reasoning', 'verbal': 'Verbal Ability',
            'python': 'Python', 'java': 'Java', 'javascript': 'JavaScript', 'sql': 'SQL',
            'algorithms': 'Algorithms', 'data structures': 'Data Structures'
        }
        for key, subject in subjects.items():
            if key.replace(' ', '') in file_name.replace(' ', ''):
                metadata["subject"] = subject
                break
        
        # Extract difficulty
        if any(term in file_name for term in ['easy', 'beginner']):
            metadata["difficulty"] = "easy"
        elif any(term in file_name for term in ['hard', 'advanced', 'expert']):
            metadata["difficulty"] = "hard"
        else:
            metadata["difficulty"] = "medium"
        
        return metadata
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def extract_text_from_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def extract_text_from_rtf(self, file_path: str) -> str:
        if not RTF_AVAILABLE:
            return f"[RTF file: {Path(file_path).name} - striprtf not available]"
        with open(file_path, 'r', encoding='utf-8') as file:
            rtf_content = file.read()
            return rtf_to_text(rtf_content)
    
    def extract_text_from_doc(self, file_path: str) -> str:
        if not DOC_AVAILABLE:
            return f"[DOC file: {Path(file_path).name} - docx2txt not available]"
        return docx2txt.process(file_path)
    
    def extract_text_from_image(self, file_path: str) -> str:
        if not OCR_AVAILABLE:
            return f"[Image file: {Path(file_path).name} - OCR not available]"
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            print(f"OCR failed for {file_path}: {e}")
            return f"[Image file: {Path(file_path).name} - OCR failed]"
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        elif file_extension == '.rtf':
            text = self.extract_text_from_rtf(file_path)
        elif file_extension == '.doc':
            text = self.extract_text_from_doc(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            text = self.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        base_metadata = self.extract_metadata(file_path, text)
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = base_metadata.copy()
            metadata.update({"chunk_id": i, "total_chunks": len(chunks)})
            
            doc = LangchainDocument(page_content=chunk, metadata=metadata)
            documents.append(doc)
        
        return documents
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        all_documents = []
        supported_extensions = ['.pdf', '.docx', '.txt', '.rtf', '.doc', '.jpg', '.jpeg', '.png', '.bmp', '.tiff']
        
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                if Path(file).suffix.lower() in supported_extensions:
                    file_path = os.path.join(root, file)
                    try:
                        documents = self.process_document(file_path)
                        all_documents.extend(documents)
                        print(f"Processed: {file_path} ({len(documents)} chunks)")
                    except Exception as e:
                        print(f"Error processing {file_path}: {str(e)}")
        
        return all_documents